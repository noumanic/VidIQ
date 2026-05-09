"""Generate the word-for-word speaker scripts for the 33-slide VidIQ deck.

Re-run with:
    ./venv/Scripts/python.exe marketing/build_speaker_scripts.py
Output:
    VidIQ_Submission/08_Speaker_Scripts.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "VidIQ_Submission" / "08_Speaker_Scripts.docx"

# Brand palette
VIOLET = RGBColor(0xA8, 0x55, 0xF7)
FUCHSIA = RGBColor(0xEC, 0x48, 0x99)
CYAN = RGBColor(0x06, 0xB6, 0xD4)
EMERALD = RGBColor(0x10, 0xB9, 0x81)
AMBER = RGBColor(0xF5, 0x9E, 0x0B)
DARK = RGBColor(0x1B, 0x14, 0x2A)
MID = RGBColor(0x4A, 0x42, 0x5C)
LOW = RGBColor(0x6B, 0x60, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, hex_rgb: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_rgb)
    tc_pr.append(shd)


def hr(doc, color="A855F7", thickness=14):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(thickness))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    p_pr.append(pbdr)


def run_set(run, *, font="Calibri", size=11, color=DARK, bold=False, italic=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_p(doc, text, *, size=11, color=DARK, bold=False, italic=False, align=None,
          space_after=4, font="Calibri"):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    run_set(r, font=font, size=size, color=color, bold=bold, italic=italic)
    return p


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def slide_card(doc, *, num, title, pillar, speaker, time, cue, script):
    """Render one slide's speaker block — meta header, visual cue, and the
    word-for-word script. Designed to fit one slide per page when printed."""
    # Top header strip — slide number + title
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cells = table.rows[0].cells
    cells[0].width = Cm(2.5)
    cells[1].width = Cm(13.5)
    cells[0].text = ""
    cells[1].text = ""
    rn = cells[0].paragraphs[0].add_run(f"#{num:02d}")
    run_set(rn, font="Calibri", size=22, color=WHITE, bold=True)
    cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(cells[0], "A855F7")
    rt = cells[1].paragraphs[0].add_run(title)
    run_set(rt, font="Calibri", size=15, color=WHITE, bold=True)
    set_cell_bg(cells[1], "1B142A")

    # Meta row — pillar · speaker · time
    meta_table = doc.add_table(rows=1, cols=3)
    for c, w in zip(meta_table.rows[0].cells, [Cm(5.5), Cm(5.5), Cm(5.0)]):
        c.width = w
    pillar_cell, speaker_cell, time_cell = meta_table.rows[0].cells

    def meta(cell, label, value, color_hex):
        cell.text = ""
        p = cell.paragraphs[0]
        rl = p.add_run(label.upper() + "  ")
        run_set(rl, size=8, color=LOW, bold=True)
        rv = p.add_run(value)
        run_set(rv, size=10, color=DARK, bold=True)
        set_cell_bg(cell, color_hex)

    meta(pillar_cell,  "Pillar",  pillar,   "F4F1F8")
    meta(speaker_cell, "Speaker", speaker,  "FCE7F3")
    meta(time_cell,    "Time",     time,    "E6F8FA")

    # Visual cue
    add_p(doc, "VISUAL CUE", size=8, color=VIOLET, bold=True, space_after=1)
    add_p(doc, cue, size=10, color=MID, italic=True, space_after=8)

    # Script
    add_p(doc, "SCRIPT — WORD FOR WORD", size=8, color=FUCHSIA, bold=True, space_after=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    r = p.add_run("“" + script + "”")
    run_set(r, size=11, color=DARK)

    # Spacer rule
    hr(doc, color="DDDDDD", thickness=8)


# ── Cover + intro ─────────────────────────────────────────────────────────


def cover(doc):
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

    add_p(doc, "SEMESTER COURSE PROJECT  ·  PRESENTATION SCRIPTS",
          size=10, color=VIOLET, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
          space_after=4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("VidIQ")
    run_set(r, size=54, color=DARK, bold=True)

    add_p(doc, "Speaker Notes & Word-for-Word Scripts — All 33 Slides",
          size=14, color=VIOLET, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    hr(doc, color="A855F7", thickness=22)

    add_p(doc, "Course",       size=10, color=LOW, bold=True)
    add_p(doc, "Digital Marketing  ·  Section CS-A  ·  Sir Maaz Zafar Cheema", size=12)

    add_p(doc, "Group members", size=10, color=LOW, bold=True, space_after=2)
    members = [
        ("22i-1653", "Insharah Aman"),
        ("21i-0416", "M. Nouman Hafeez"),
        ("21i-0484", "Shayan Khan"),
        ("21i-2507", "Muhammad Zain"),
        ("22i-1200", "Farhan Ahmed"),
    ]
    for rno, name in members:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        rrn = p.add_run(rno + "   ")
        run_set(rrn, font="Consolas", size=11, color=CYAN, bold=True)
        rnm = p.add_run(name)
        run_set(rnm, size=11, color=DARK)

    add_p(doc, "Total presentation time", size=10, color=LOW, bold=True,
          space_after=2)
    add_p(doc, "12-15 minutes presented + 5-10 minutes Q&A  ·  total target: 20 minutes",
          size=11)

    page_break(doc)


def speaker_assignment(doc):
    add_p(doc, "Speaker Assignment", size=18, color=VIOLET, bold=True, space_after=4)
    hr(doc, color="A855F7", thickness=18)
    add_p(doc,
          "Slides are assigned to one of five speakers based on pillar ownership. Each speaker "
          "rehearses their own slides and is the primary defender for those rubric points "
          "during Q&A. The full deck is 33 slides — 19 are presented live (slides 1-19), "
          "and slides 20-33 are appendix/deep-dive backup that can be pulled up on demand.",
          size=11, space_after=12)

    table = doc.add_table(rows=6, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["Speaker", "Pillar Ownership", "Live Slides", "Backup Slides (on demand)"]
    for j, h in enumerate(headers):
        c = table.rows[0].cells[j]
        c.text = ""
        rh = c.paragraphs[0].add_run(h)
        run_set(rh, size=10, color=WHITE, bold=True)
        set_cell_bg(c, "1B142A")

    rows = [
        ("M. Nouman Hafeez (21i-0416)",
         "Lead · Product",
         "1, 3, 4, 19, finale",
         "20"),
        ("Insharah Aman (22i-1653)",
         "Pillar 1 · Brand & Video Ad",
         "5, 6",
         "21, 22, 32, 33"),
        ("Shayan Khan (21i-0484)",
         "Pillar 2 · Social & Meta Ads",
         "7, 8, 9",
         "23, 24, 25"),
        ("Muhammad Zain (21i-2507)",
         "Pillar 4 · SEO & Google Ads",
         "2, 10, 11, 12",
         "26, 27, 31"),
        ("Farhan Ahmed (22i-1200)",
         "Pillar 5 · Budget · Competitive · KPIs",
         "13, 14, 15, 16, 17, 18",
         "28, 29, 30"),
    ]
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            c = table.rows[i].cells[j]
            c.text = ""
            rr = c.paragraphs[0].add_run(val)
            run_set(rr, size=10, color=DARK, bold=(j == 0))
            set_cell_bg(c, "F4F1F8" if i % 2 == 1 else "FFFFFF")

    add_p(doc, "Pacing target (12 minutes presented)", size=14, color=VIOLET,
          bold=True, space_after=4)
    pacing = [
        ("0:00 - 0:30",  "Slide 1",      "Cover · hook"),
        ("0:30 - 1:30",  "Slides 2-3",   "Problem + Product (live demo if time)"),
        ("1:30 - 2:30",  "Slide 4",      "Pipeline architecture"),
        ("2:30 - 4:00",  "Slides 5-6",   "Brand + 30-second ad (play the video)"),
        ("4:00 - 6:00",  "Slides 7-9",   "Social calendar + Meta Ads + Conversation"),
        ("6:00 - 8:30",  "Slides 10-12", "Keywords + On-page SEO + Google Ads"),
        ("8:30 - 10:30", "Slides 13-15", "Budget (PKR) + Competitive + Social-comp"),
        ("10:30 - 11:30","Slides 16-18", "SWOT + KPI tracker + Scorecard"),
        ("11:30 - 12:00","Slide 19",     "Q&A backup intro · open the floor"),
        ("12:00+",        "Slides 20-33", "Appendix backup — pulled up only on Q&A"),
    ]
    table = doc.add_table(rows=len(pacing) + 1, cols=3)
    headers = ["Window", "Slides", "Activity"]
    for j, h in enumerate(headers):
        c = table.rows[0].cells[j]
        c.text = ""
        rh = c.paragraphs[0].add_run(h)
        run_set(rh, size=10, color=WHITE, bold=True)
        set_cell_bg(c, "1B142A")
    for i, row in enumerate(pacing, start=1):
        for j, val in enumerate(row):
            c = table.rows[i].cells[j]
            c.text = ""
            rr = c.paragraphs[0].add_run(val)
            run_set(rr, font=("Consolas" if j == 0 else "Calibri"),
                    size=10, color=DARK, bold=(j == 1))
            set_cell_bg(c, "F4F1F8" if i % 2 == 1 else "FFFFFF")

    page_break(doc)


# ── Slide-by-slide scripts ────────────────────────────────────────────────


SLIDES = [
    # (num, title, pillar, speaker, time, visual cue, script)
    (1, "VidIQ — AI Video Intelligence (Cover)",
     "Cover", "M. Nouman Hafeez", "0:00 - 0:30",
     "Cover slide on screen — VidIQ wordmark · Watch less. Learn more. · live URL pill at the bottom.",
     "Good morning, sir, and good morning to the panel. We are Group VidIQ, and over the next "
     "twelve minutes we will walk you through our digital-marketing strategy and execution for "
     "VidIQ — an AI video-intelligence platform we built from scratch. The product is live at "
     "vidiq-two.vercel.app, the brand is launched on Facebook and Instagram, and every campaign "
     "you will see today exists in a real Ads Manager account. Let me hand over straight to the "
     "problem we set out to solve."),

    (2, "The job-to-be-done — long-form video is unread",
     "Problem", "Muhammad Zain", "0:30 - 1:00",
     "Stat slide — 500 hours of video uploaded to YouTube every minute · average watch time per session: 17 minutes.",
     "Five hundred hours of video is uploaded to YouTube every single minute. The average user "
     "watches just seventeen minutes of it per session. That gap — between what is published "
     "and what is actually consumed — is the job-to-be-done for VidIQ. Students miss lectures, "
     "traders miss earnings calls, professionals miss training they paid for. We built a tool "
     "that turns any video, recorded or live, into a structured summary in under a minute."),

    (3, "The product — paste a URL, get intelligence",
     "Pillar 3", "M. Nouman Hafeez", "1:00 - 1:30",
     "Live screen-record loop of vidiq-two.vercel.app — a YouTube URL gets pasted, summary tabs render, citations appear.",
     "This is the product running on the live URL. The user pastes a YouTube link, hits "
     "Analyze, and within seconds we return a faithful summary, time-stamped chapters, detected "
     "events, and a chat interface where every answer cites the exact moment in the video. "
     "Click the citation, and the embedded player seeks to that timestamp. That citation-grounded "
     "behavior is the trust moat — and it is what differentiates us from generic AI summarisers."),

    (4, "Multimodal pipeline — speech + vision + LLM",
     "Pillar 3", "M. Nouman Hafeez", "1:30 - 2:30",
     "Architecture diagram — frontend on Vercel rewrites to FastAPI on Hugging Face Space; pipeline shows transcript + keyframes + LLM map-reduce.",
     "Architecturally, we run a multimodal pipeline. Speech-to-text uses YouTube native "
     "transcripts when available and falls back to faster-whisper running locally. Vision runs "
     "OpenCV scene-change keyframe extraction, and a vision LLM captions every keyframe. A "
     "map-reduce LLM pipeline then fuses all three streams into one summary. The backend is "
     "provider-agnostic — Gemini, Groq, OpenAI, or stub — with automatic failover, so we are "
     "resilient to free-tier quota changes. Frontend is Next.js on Vercel, backend is FastAPI "
     "on Hugging Face Spaces. Both are publicly deployed and you can hit them right now."),

    (5, "Pillar 1 — Brand identity",
     "Pillar 1", "Insharah Aman", "2:30 - 3:15",
     "Four-up panel — logo · violet/fuchsia/cyan palette · Plus Jakarta Sans + Inter sample · voice slider.",
     "Now to Pillar One — branding. Our mark is a stylised play-head that reads as the letter "
     "V — communicating video without a literal play triangle. The wordmark splits Vid in solid "
     "sans from IQ in a violet-to-fuchsia gradient, putting the intelligence claim front and "
     "centre. Violet was a deliberate choice. Blue is the SaaS default — Stripe, Twitch, "
     "Discord — and we wanted the trust association of blue with the creativity of red, which "
     "violet sits exactly between. Typography is Plus Jakarta Sans for display and Inter for "
     "body — both humanist geometric, both modern. The voice is approachable but credible — "
     "restrained, never zany."),

    (6, "Pillar 1 — The 30-second ad",
     "Pillar 1", "Insharah Aman", "3:15 - 4:00",
     "Embedded 30-second hero video plays — six-beat story-arc bar across the bottom.",
     "And here is the 30-second hero ad we cut for the launch flight. The structure is a "
     "six-beat narrative arc: hard hook in the first three seconds, promise at five, product "
     "proof at eight, our differentiator — live streams — at fifteen, social proof at "
     "twenty-two, and a call-to-action loop in the last three seconds. We mastered this once "
     "in nine-by-sixteen and exported four cutdowns — fifteen seconds, nine seconds, six "
     "seconds, and one-by-one square — for placement-specific delivery on Reels, Stories, "
     "YouTube In-Stream, and feed."),

    (7, "Pillar 2 — Social calendar · FB + IG",
     "Pillar 2", "Shayan Khan", "4:00 - 4:45",
     "Business Suite Planner screenshot · 28 posts queued · FB page · IG profile screenshots on the right.",
     "Pillar Two — social media. Both identities are live. Our Facebook page is at "
     "facebook.com/profile.php with the ID shown on the slide, and our Instagram is "
     "@vidiq_official. The 14-day calendar is built and queued inside Meta Business Suite — "
     "twenty-eight posts in total, two per day, alternating Facebook and Instagram. The format "
     "mix covers all four required formats: Reels, Carousels, Single images, and Stories. Each "
     "post carries between eight and twelve hashtags split across brand, category, and audience "
     "tags. The cadence — seven posts per week per channel — is more aggressive than either "
     "competitor, NoteGPT or Eightify."),

    (8, "Pillar 2 — Meta Ads · 4 ad sets to Review",
     "Pillar 2", "Shayan Khan", "4:45 - 5:30",
     "Two-by-two grid of Ads Manager Review screenshots — Students · Creators · Knowledge · Retargeting.",
     "Our Meta flight runs four ad sets, each targeting a distinct audience cluster. Set one is "
     "students aged eighteen to twenty-four, with a Reel that promises two hours of lecture "
     "compressed into five minutes of notes. Set two is creators aged twenty-two to forty, with "
     "a carousel about saving four hours per week of research. Set three is knowledge workers "
     "aged twenty-eight to forty-five, with a Reel about turning earnings calls into action "
     "items. Set four is retargeting site visitors and Instagram engagers. All four are built "
     "end-to-end in Ads Manager and screenshotted at the Review screen — they are ready to "
     "launch but stay in Draft for the showcase. Total Meta budget is ninety-five thousand "
     "rupees, split evenly at twenty-three thousand seven hundred fifty per set."),

    (9, "Pillar 3 — Always-on conversation",
     "Pillar 2 + 3", "Shayan Khan", "5:30 - 6:00",
     "Four screenshots — welcome message · auto-FAQ · saved replies · auto-away.",
     "Every paid impression that converts to a profile visit lands in a conversation surface — "
     "Messenger, Instagram DM, or our in-app chat. We configured four automation primitives in "
     "Meta Business Suite. A welcome message fires on the first DM and offers the demo URL. An "
     "auto-FAQ responder handles the six most-asked intents — pricing, how it works, supported "
     "videos, privacy, live streams, and refunds. Saved replies give human agents an eight-pack "
     "library for fast responses. And an auto-away message sets expectations outside business "
     "hours. This automation cuts median first-response from hours to seconds."),

    (10, "Pillar 4 — Keyword strategy · 18 KWs",
     "Pillar 4", "Muhammad Zain", "6:00 - 6:45",
     "Three-tier funnel · 5 head terms · 4 long-tail · 3 niche · MSV and KD per keyword.",
     "Pillar Four is search. We researched eighteen keywords and split them into three tiers. "
     "Tier one is head terms like ai video summarizer with twenty-seven thousand monthly "
     "searches but a keyword difficulty of forty-eight — saturated. Tier two is long-tail like "
     "summarize a two-hour youtube video with keyword difficulty under fifteen — rankable in "
     "two to three months. Tier three is niche, vertical-domain queries like medical lecture "
     "video summarizer where there is zero direct competition. Our strategy is to bid head "
     "terms via Google Ads, and rank long-tail and niche organically. We do not waste "
     "domain-authority-one effort on terms NoteGPT already owns."),

    (11, "Pillar 4 — On-page SEO · 8 techniques",
     "Pillar 4", "Muhammad Zain", "6:45 - 7:30",
     "Eight technique cards in a four-by-two grid · proofs row at the bottom.",
     "We applied eight on-page SEO techniques across the live site. Per-page meta tags with "
     "Open Graph and Twitter cards. Descriptive alt text on every image — auto-generated from "
     "our vision captions. Strict header hierarchy from H1 to H3, validated with Wave and axe "
     "DevTools. JSON-LD structured data — Organization, WebSite, and SoftwareApplication "
     "schemas — for rich-result eligibility. A live sitemap and robots.txt, both submitted to "
     "Google Search Console. Hub-and-spoke internal linking from the home page out to "
     "/analyze, /live, and /library. Core Web Vitals optimisation through next/image, "
     "lazy-loading, and Vercel Edge caching — we are well under the LCP and CLS thresholds. "
     "And mobile-first responsive plus HTTPS via Vercel. Lighthouse SEO scores one hundred. "
     "The proofs are screenshot-able — sitemap dot xml, robots dot txt, view-source for the "
     "JSON-LD, and the Search Console verified property."),

    (12, "Pillar 4 — Google Ads · 3 campaigns to Review",
     "Pillar 4", "Muhammad Zain", "7:30 - 8:15",
     "Three Review screenshots — Search · YouTube · PMax — bidding rules at the bottom.",
     "We built three Google Ads campaigns in parallel with the Meta flight, totalling seventy "
     "thousand rupees. Search Brand Cluster runs four ad groups bidding on our brand terms plus "
     "competitor terms plus high-intent long-tail — twenty-three thousand one hundred rupees. "
     "YouTube In-Stream uses thirty-second skippable ads and six-second bumpers at a target CPM "
     "— eighteen thousand nine hundred rupees. Performance Max uses Maximise Conversion Value "
     "across three asset groups — twenty-eight thousand rupees, our biggest single line because "
     "PMax compounds across YouTube, Display, and Search simultaneously. The bidding ladder is "
     "Maximise Conversions for the first thirty conversions, then we switch to a Target CPA cap "
     "of four hundred twenty rupees, which is roughly one dollar fifty. PMax is gated on "
     "thirty conversions in seven days — if it does not unlock we reallocate to Search."),

    (13, "Pillar 5 — Budget · ₨ 250,000 PKR",
     "Pillar 5", "Farhan Ahmed", "8:15 - 9:00",
     "Six-line budget card · PKR primary · USD secondary · rationale callout · showcase actual ₨ 0.",
     "Pillar Five — budget. The launch flight is sized at two hundred and fifty thousand "
     "Pakistani rupees, which is approximately eight hundred and ninety US dollars at "
     "today's exchange rate. We budget in rupees because we operate from Pakistan and pay "
     "Meta and Google in rupees. The split: ninety-five thousand to Meta Ads, seventy thousand "
     "to Google Ads, forty thousand to creative production, twenty thousand to influencer "
     "seeding, seventeen and a half thousand to tools and SaaS, seventy-five hundred to "
     "contingency. Why two hundred and fifty thousand specifically? Pakistani SaaS launch "
     "budgets cluster between one hundred fifty and four hundred thousand per month, so this "
     "is mid-range and aggressive for a fourteen-day flight. Pakistani Meta CPM averages "
     "eighty to two hundred and fifty rupees versus five to fifteen US dollars in the West — "
     "five to seven times more reach per rupee. Two fifty K buys roughly eighty thousand to "
     "one hundred and twenty thousand impressions, enough for the pixel to leave learning. "
     "Showcase actual is zero rupees — every campaign sits in Draft."),

    (14, "Pillar 5 — Competitive · NoteGPT vs Eightify",
     "Pillar 5", "Farhan Ahmed", "9:00 - 9:30",
     "Three-column table · pricing · features · live streams · multimodal · languages.",
     "Two direct competitors. NoteGPT charges nine dollars per month with ten free summaries. "
     "Eightify charges eight ninety-nine with five free summaries. Both handle YouTube. "
     "Neither handles live streams. Neither is multimodal in the sense that we are — we "
     "process vision, audio, and text together. NoteGPT supports twelve languages, Eightify "
     "five — we support eleven plus right-to-left. Our domain modes for medical, legal, and "
     "trading do not exist on either competitor. The price-point we will land on is also "
     "nine dollars per month, but with a more generous free tier — sub-minute runs without "
     "monthly count limits."),

    (15, "Pillar 5 — Social presence comparison",
     "Pillar 5", "Farhan Ahmed", "9:30 - 10:00",
     "Two side-by-side tables — Facebook and Instagram · VidIQ vs NoteGPT vs Eightify.",
     "Socially we are behind on raw follower count — both competitors had a multi-year head "
     "start. NoteGPT sits at roughly three thousand one hundred Facebook followers and five "
     "thousand eight hundred on Instagram. Eightify is at six hundred twenty and one thousand "
     "nine hundred. We are launching from zero. But our planned cadence — seven posts per "
     "week per channel — is more aggressive than either, and our hashtag depth — eight to "
     "twelve per Instagram post — beats NoteGPT's six and Eightify's three. We will close the "
     "follower gap with the Meta ad spend you saw two slides ago."),

    (16, "Pillar 5 — SWOT analysis",
     "Pillar 5", "Farhan Ahmed", "10:00 - 10:30",
     "Two-by-two SWOT grid · Strengths · Weaknesses · Opportunities · Threats.",
     "Quick SWOT. Strengths: multimodal AI, live-stream pipeline, provider-agnostic backend, "
     "domain modes, and citation grounding. Weaknesses: zero brand awareness today, "
     "domain-authority of one, brand-name collision with the unrelated vidiq.com, no native "
     "mobile yet. Opportunities: long-tail SEO is rankable in months not years, vertical "
     "niches are first-mover territory, and Pakistani CPM arbitrage gives us five-to-seven-X "
     "more impressions per rupee. Threats: big-tech could ship native summaries, LLM API "
     "pricing could compress margins, and YouTube terms-of-service could change around "
     "transcript scraping."),

    (17, "KPI tracker · all 18 graded KPIs",
     "Pillar 5", "Farhan Ahmed", "10:30 - 11:00",
     "Full table of 18 KPIs · pillar · target · status · evidence file.",
     "We track all eighteen graded KPIs in a live dashboard. Brand consistency: five out of "
     "five — eight surfaces audited green. Social calendar: fourteen out of fourteen days "
     "queued. Format mix: four out of four formats. Both social identities live. Four Meta "
     "ad sets built to Review. Welcome and auto-FAQ configured. Saved replies and auto-away "
     "configured. Eighteen keywords across three tiers. Eight on-page SEO techniques applied. "
     "Search Console verified, sitemap submitted. Three Google Ads campaigns built. Bidding "
     "strategy with CPA cap and PMax gating documented. Total budget in PKR with full "
     "distribution. Two-competitor analysis. Every KPI has a verifiable artefact — a URL or "
     "a screenshot — and they are all in the submission package."),

    (18, "KPI scorecard · pillar-level summary",
     "Pillar 5", "Farhan Ahmed", "11:00 - 11:30",
     "Pillar-level scorecard · each pillar marked Strong · 15 / 15 rubric points.",
     "Rolled up to the pillar level: branding fifteen out of fifteen on the four KPIs, social "
     "media four of four, conversation three of three, SEO and ads five of five, and "
     "competitive analysis with budget two of two. Self-assessment is fifteen out of fifteen "
     "across the five pillar deliverables. The remaining ten marks come from this presentation "
     "and the Q&A defense — which is what you are about to grade now."),

    (19, "Q&A — open the floor",
     "Defense", "M. Nouman Hafeez", "11:30 - 12:00",
     "Q&A backup slide · live demo URL pinned to the corner.",
     "That brings us to the close. Everything you have seen in the last twelve minutes is on "
     "the live URL — vidiq-two.vercel.app. The deck, the report, the Excel KPI sheet, every "
     "screenshot, the video ad, and our brand assets are in the submission folder we sent you. "
     "Both social profiles are live and posting. We are ready for your questions — and we have "
     "appendix slides covering deeper dives on the pipeline, the brand system, the ad funnel, "
     "the conversation flow, the keyword research, the budget phasing, the moats, and the "
     "roadmap, which we can pull up on demand. Sir, the floor is yours."),

    # ── Appendix — backup slides 20-33 ──
    (20, "Strategy overview · five pillars · positioning · GTM thesis",
     "Appendix", "M. Nouman Hafeez", "appendix · 30 s on demand",
     "Five-pillar strategy overview · positioning vs NoteGPT/Eightify · go-to-market thesis.",
     "Our overall go-to-market thesis is that we acquire users by ranking for long-tail "
     "informational queries via SEO, capture head-term demand via Google Ads at the moment of "
     "intent, and build a brand with always-on social presence on Facebook and Instagram. The "
     "five pillars in this slide map exactly to the rubric — branding, social, conversation, "
     "search and ads, and the budget-and-competitive layer."),

    (21, "Brand identity · palette + typography + voice",
     "Pillar 1 · Appendix", "Insharah Aman", "appendix · 30 s on demand",
     "Color tokens with hex + role · typography with three typefaces · voice slider.",
     "Deeper dive on brand. Six tokens — background dark, primary violet, fuchsia accent, "
     "cyan accent, emerald success, amber warning. Three typefaces — Plus Jakarta Sans for "
     "display, Inter for body, JetBrains Mono for code and timestamps. The voice slider sits "
     "at approachable-but-credible, restrained-not-zany, and modern-indie."),

    (22, "Video ad campaign · 6-beat script + 4 cutdowns",
     "Pillar 1 · Appendix", "Insharah Aman", "appendix · 45 s on demand",
     "Beat-by-beat table of the 30s ad · timecodes · visuals · voiceover · 4 cutdown specs.",
     "If you want to see the ad's anatomy, we have the full beat-by-beat script. Hook at "
     "zero seconds, promise at three, proof at eight, differentiator at fifteen, social proof "
     "at twenty-two, call-to-action at twenty-seven. The four cutdowns are fifteen-second "
     "TikTok-native, nine-second pre-roll, six-second bumper, and one-by-one square. All "
     "exported from the same master timeline so the look is consistent."),

    (23, "Content calendar · 14-day pillars and format mix",
     "Pillar 2 · Appendix", "Shayan Khan", "appendix · 30 s on demand",
     "14-day calendar grid · content pillars · format mix table · cadence breakdown.",
     "The fourteen-day content calendar runs across four pillars — Product, Education, Brand, "
     "and Community. Format cadence: three Reels per week, two Carousels, one single image, "
     "and daily Stories. Every post is scheduled in Business Suite and queued."),

    (24, "Meta Ads · 4 ad sets deep dive",
     "Pillar 2 · Appendix", "Shayan Khan", "appendix · 45 s on demand",
     "Each ad set with audience definition · creative type · KPI per set.",
     "Per-ad-set deep dive. Set one Students wants cost-per-result under four hundred twenty "
     "rupees, Reel-led. Set two Creators wants higher CTR — carousel format. Set three "
     "Knowledge Workers is the highest-LTV target so we accept a higher cost-per-result. Set "
     "four Retargeting is our cheapest acquisition because the audience is warm — site "
     "visitors and Instagram engagers."),

    (25, "Conversation layer · welcome · FAQs · saved replies",
     "Pillar 3 · Appendix", "Shayan Khan", "appendix · 30 s on demand",
     "Detailed walkthrough of welcome · FAQs · saved replies · auto-away with screenshots.",
     "Conversation deep dive. The welcome message is two lines plus the demo URL. The auto-"
     "FAQ has six trained intents — pricing, how it works, supported videos, privacy, live "
     "streams, refunds. The saved-replies library has eight entries for fast human-agent "
     "answers. Auto-away is bound to working hours and sets expectations on response window."),

    (26, "Keyword + on-page SEO deep dive",
     "Pillar 4 · Appendix", "Muhammad Zain", "appendix · 45 s on demand",
     "Keyword funnel by tier · on-page audit grid · 18 KWs by KD and MSV.",
     "Eighteen keywords across three tiers — five head, ten long-tail, three niche. On-page "
     "audit is four elements: meta tags, alt text, header hierarchy, and keyword placement. "
     "All four are green on the live site. Strategy: own long-tail and niche organically, "
     "bid head terms via Google Ads."),

    (27, "Google Ads · 3 campaigns deep dive",
     "Pillar 4 · Appendix", "Muhammad Zain", "appendix · 45 s on demand",
     "Each Google campaign with bidding rules · audience signals · creative slots.",
     "Per-campaign Google Ads deep dive. Search Brand Cluster has four ad groups — brand, "
     "competitor, generic, and long-tail — each with three ads in rotation. YouTube In-Stream "
     "has both thirty-second skippable and six-second bumper variants. Performance Max runs "
     "three asset groups corresponding to our three highest-intent audience signals."),

    (28, "Budget distribution · phasing · paid-media split",
     "Pillar 5 · Appendix", "Farhan Ahmed", "appendix · 45 s on demand",
     "Detailed PKR breakdown · Meta + Google split · 4-phase spend timeline.",
     "Detailed budget phasing. Days one through seven are Learn — highest-volume bid, "
     "landing-page-views objective, cheapest data while the pixel learns. Day five we open a "
     "creative-swap fund — ten percent of Meta budget reserved. Days eight through fourteen "
     "switch to Conversions with a CPR cap of four hundred and twenty rupees. PMax unlocks if "
     "we bank thirty conversions. Day fourteen is the negative-keyword sweep, screenshot "
     "reports, publish learnings."),

    (29, "Conversion tracking · Pixel · CAPI · gtag · GA4",
     "Pillar 2+4 · Appendix", "Farhan Ahmed", "appendix · 30 s on demand",
     "Tracking architecture diagram · same event names across stack · KPI targets table.",
     "Tracking architecture: Meta Pixel browser-side plus the Conversions API server-side from "
     "our FastAPI backend, both firing the same event names so iOS-fourteen breakage does not "
     "blind us. Google gtag site-wide. GA4 linked to Search Console. Event taxonomy: view "
     "demo, start analysis, finish analysis, open chat, sign up. Same names across the stack "
     "to avoid dashboard drift."),

    (30, "Moats + roadmap · 12-month outlook",
     "Pillar 5 · Appendix", "Farhan Ahmed", "appendix · 30 s on demand",
     "5 moats card · 8-phase roadmap timeline.",
     "Five moats we are deliberately building — citations, live streams, vertical domain "
     "modes, provider resilience, and an open-source-able codebase. Twelve-month roadmap goes "
     "from launch flight today, to first thousand users, to native mobile, to vertical "
     "bundles, to browser extension, to team accounts, to embedding API, to Series A "
     "readiness."),

    (31, "SEO evidence trail · GSC · sitemap · Lighthouse",
     "Pillar 4 · Appendix", "Muhammad Zain", "appendix · 30 s on demand",
     "5 evidence screenshots · property verified · sitemap submitted · URL inspection · PageSpeed mobile + desktop.",
     "Live SEO evidence. Search Console property verified by HTML-tag method. Sitemap "
     "submitted with four URLs discovered. URL inspection requesting indexing for home, "
     "/analyze, and /live. PageSpeed Insights audited on both mobile and desktop. Every "
     "screenshot was captured on May sixth twenty twenty-six and is stored in the submission "
     "folder under SEO evidence."),

    (32, "Logo system · 3 lockups · brand rules",
     "Pillar 1 · Appendix", "Insharah Aman", "appendix · 30 s on demand",
     "Three logo lockups in PNG and SVG · square dark · square light · horizontal wordmark · clear-space rules.",
     "Three logo lockups maintained as the single source of truth. Square mark on dark for "
     "favicons and app icons. Square mark on light for OG cards and print. Horizontal "
     "wordmark for ad creatives and video end-cards. Clear-space rule is one-times-mark-height "
     "on all four sides. Minimum sizes are twenty-four pixels for digital favicon, forty for "
     "UI navigation, ninety-six for print."),

    (33, "Logo applied · 7 surfaces audited",
     "Pillar 1 · Appendix", "Insharah Aman", "appendix · 30 s on demand · CLOSE",
     "Seven applied surfaces · favicon · top nav · hero splash · OG · JSON-LD · slide footer · FB cover · IG profile.",
     "Final slide. The mark is applied across seven live surfaces — favicon, top nav, hero "
     "splash, Open Graph card, JSON-LD logo URL, every slide footer in this deck, and both "
     "social profiles. The brand-consistency KPI scores five out of five. Thank you, sir. "
     "Floor is yours."),
]


def build_slides(doc):
    add_p(doc, "Slide-by-slide scripts", size=18, color=VIOLET, bold=True, space_after=4)
    hr(doc, color="A855F7", thickness=18)
    add_p(doc,
          "Each slide below shows the slide number, title, owning pillar, the speaker, the "
          "time-window, the visual cue on screen, and the word-for-word script. Read in your "
          "natural pace — the timings assume roughly 130 words per minute. Slides 1-19 are the "
          "live walkthrough; slides 20-33 are appendix backup pulled up only if Q&A demands.",
          size=11, space_after=10)

    for tup in SLIDES:
        slide_card(doc,
                    num=tup[0], title=tup[1], pillar=tup[2], speaker=tup[3],
                    time=tup[4], cue=tup[5], script=tup[6])

    page_break(doc)


# ── Q&A defense ───────────────────────────────────────────────────────────


def qa_defense(doc):
    add_p(doc, "Q&A defense — pre-built answers",
          size=18, color=VIOLET, bold=True, space_after=4)
    hr(doc, color="A855F7", thickness=18)
    add_p(doc,
          "Likely instructor questions and one-sentence pre-built answers. Speaker assignments "
          "in brackets — that person owns the answer if asked.",
          size=11, space_after=10)

    qas = [
        ("Why violet, not blue?  [Insharah]",
         "Blue is the SaaS default — Stripe, Linear, Twitch — violet retains blue's trust association "
         "while signalling creativity and intelligence; it is the colour of insight."),
        ("Why both Meta Pixel and CAPI?  [Farhan]",
         "iOS 14 broke client-side pixel reliability — server-side CAPI fires from our FastAPI backend "
         "on the same event names so attribution stays measurable on iOS users."),
        ("Why long-tail keywords first?  [Muhammad Zain]",
         "Domain authority of 1 — head terms are unrankable in 14 days. Long-tail with KD under 15 is "
         "achievable with one optimised page each in 2-3 months."),
        ("How do you compete with YouTube's own AI summary?  [Nouman]",
         "Three durable moats — live-stream support, multimodal grounding with timestamp citations, and "
         "vertical-domain modes for medical, legal, and trading. YouTube's summary is generic and not citation-grounded."),
        ("What is the biggest risk?  [Nouman]",
         "Free-tier API tightening from Gemini or Groq. Our provider abstraction in services/llm.py "
         "auto-fails over Gemini → Groq → OpenAI → stub, so a single revocation does not break us."),
        ("Why ₨ 250,000 specifically — could you do it cheaper?  [Farhan]",
         "Yes, technically — but ₨ 250 K is the threshold where the Meta pixel banks enough conversions "
         "to leave the Learning phase before D8 and converge on a stable CPA. Below ₨ 150 K we waste budget on noisy data."),
        ("Why not split the budget more toward Google?  [Farhan]",
         "Google captures intent — but at DA-1 we have no Quality Score yet, so CPCs are punitive. "
         "Meta's cheaper CPM lets us build pixel data, retarget, and warm the audience before Google takes over in week 2."),
        ("Plagiarism — prove this is original  [Nouman]",
         "Every screenshot is from a live, working account we own — Vercel, Hugging Face, Search "
         "Console, Meta Business Suite, Google Ads. Source URLs are in 07_Links.txt and the GitHub repo. "
         "The copy was written by us; any data points cite their public source."),
        ("Why Hugging Face for the backend?  [Nouman]",
         "Free-tier-first — HF Spaces gives us a Docker container with CPU-basic at no cost, including "
         "ffmpeg and faster-whisper. Vercel free tier hosts the frontend. Total infrastructure cost: zero rupees."),
        ("Pakistan-specific reach claims — proof?  [Farhan]",
         "Pakistani Meta CPM averages ₨ 80-250 per 1000 impressions according to Meta's own benchmarking "
         "tool inside Ads Manager (visible on the Audience-size estimator). Western CPM is publicly "
         "documented at $5-15 by Meta. Ratio: 5-7×."),
        ("How will you measure success after launch?  [Farhan]",
         "GA4 events — start_analysis, finish_analysis, sign_up — feed both Google Ads conversion "
         "tracking and Meta CAPI. KPI targets: cost-per-result ≤ ₨ 420, CTR ≥ 1.2% on Meta and ≥ 4% on "
         "Google Search, 30 demo starts/day, 8 sign-ups/day."),
        ("Have you actually run any of these ads?  [Shayan]",
         "No — and that is intentional per the brief. The rubric grades planned campaigns built to the "
         "Review screen, not paid reach. Every campaign sits in Draft, ready to launch, screenshotted at Review."),
    ]

    table = doc.add_table(rows=len(qas) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    h_cells = table.rows[0].cells
    for j, h in enumerate(["Question (& assigned defender)", "One-sentence answer"]):
        h_cells[j].text = ""
        rh = h_cells[j].paragraphs[0].add_run(h)
        run_set(rh, size=10, color=WHITE, bold=True)
        set_cell_bg(h_cells[j], "1B142A")
    for j, w in enumerate([Cm(6.5), Cm(9.5)]):
        for c in table.columns[j].cells:
            c.width = w
    for i, (q, a) in enumerate(qas, start=1):
        cells = table.rows[i].cells
        cells[0].text = ""
        cells[1].text = ""
        rq = cells[0].paragraphs[0].add_run(q)
        run_set(rq, size=10, color=DARK, bold=True)
        ra = cells[1].paragraphs[0].add_run(a)
        run_set(ra, size=10, color=DARK)
        set_cell_bg(cells[0], "F4F1F8" if i % 2 == 1 else "FFFFFF")
        set_cell_bg(cells[1], "F4F1F8" if i % 2 == 1 else "FFFFFF")

    add_p(doc, "Universal closing line (any speaker, after a long question)",
          size=11, color=VIOLET, bold=True, space_after=2)
    add_p(doc,
          "“That is documented in the report — Section [N] — and the supporting screenshot is in "
          "the submission folder under 05_Screenshots/. Happy to walk through it after if useful.”",
          size=11, italic=True, color=DARK)


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    cover(doc)
    speaker_assignment(doc)
    build_slides(doc)
    qa_defense(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUT)
    except PermissionError:
        i = 2
        while OUT.with_name(f"08_Speaker_Scripts_v{i}.docx").exists():
            i += 1
        out = OUT.with_name(f"08_Speaker_Scripts_v{i}.docx")
        doc.save(out)
        print(f"[warn] primary file locked - wrote {out.name}")
        return
    print(f"[ok] wrote {OUT}")


if __name__ == "__main__":
    main()
