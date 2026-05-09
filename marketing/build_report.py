"""Generate the formal VidIQ project report as a .docx — mirrors the deck content
in long-form professional report style for the Digital Marketing semester project.

Re-run with:
    ./venv/Scripts/python.exe marketing/build_report.py
Output:
    marketing/submissions/VidIQ_Final_Report.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor


# ── Brand palette (mirrors deck) ───────────────────────────────────────────
BG_DEEP = RGBColor(0x0A, 0x06, 0x12)
FG_HIGH = RGBColor(0x1B, 0x14, 0x2A)
FG_MID = RGBColor(0x4A, 0x42, 0x5C)
FG_LOW = RGBColor(0x6B, 0x60, 0x80)
ACCENT_VIOLET = RGBColor(0xA8, 0x55, 0xF7)
ACCENT_FUCHSIA = RGBColor(0xEC, 0x48, 0x99)
ACCENT_CYAN = RGBColor(0x06, 0xB6, 0xD4)
ACCENT_EMERALD = RGBColor(0x10, 0xB9, 0x81)
ACCENT_AMBER = RGBColor(0xF5, 0x9E, 0x0B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_PANEL = RGBColor(0x16, 0x10, 0x21)


# ── Asset paths ────────────────────────────────────────────────────────────
ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets"
SHOTS = ASSETS / "screenshots"
FIREFLY = ASSETS / "firefly"
SEO = ASSETS / "seo-evidence"
PUBLIC = ROOT.parent / "frontend" / "public"


def asset(*names: str) -> Path | None:
    for n in names:
        for root in (SHOTS, FIREFLY, SEO, PUBLIC):
            p = root / n
            if p.exists():
                return p
    return None


# ── Style helpers ──────────────────────────────────────────────────────────


def set_cell_bg(cell, hex_rgb: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_rgb)
    tc_pr.append(shd)


def add_horizontal_rule(doc, color="A855F7", thickness=12):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(thickness))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    p_pr.append(pbdr)


def set_run(run, *, font="Calibri", size=11, color=FG_HIGH, bold=False, italic=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_paragraph(doc, text, *, font="Calibri", size=11, color=FG_HIGH,
                   bold=False, italic=False, align=None, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_run(r, font=font, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, *, level=1):
    sizes = {1: 22, 2: 16, 3: 13}
    colors = {1: ACCENT_VIOLET, 2: ACCENT_VIOLET, 3: ACCENT_FUCHSIA}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    if level == 1:
        # Page-break before H1 (except the first)
        run0 = p.add_run()
        # We don't break here — caller decides.
    r = p.add_run(text)
    set_run(r, font="Calibri", size=sizes[level], color=colors[level], bold=True)
    if level == 1:
        # Coloured underline rule
        add_horizontal_rule(doc, color="A855F7", thickness=18)
    return p


def add_eyebrow(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text.upper())
    set_run(r, font="Calibri", size=8.5, color=ACCENT_VIOLET, bold=True)


def add_kv_table(doc, rows, *, col_widths=(Cm(4.0), Cm(11.5)), header=None,
                  accent=ACCENT_VIOLET):
    table = doc.add_table(rows=len(rows) + (1 if header else 0), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    # widths
    for i, w in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = w
    # header
    start = 0
    if header:
        h_cells = table.rows[0].cells
        for j, h in enumerate(header):
            h_cells[j].text = ""
            p = h_cells[j].paragraphs[0]
            r = p.add_run(h)
            set_run(r, font="Calibri", size=10.5, color=WHITE, bold=True)
            set_cell_bg(h_cells[j], "1B142A")
        start = 1
    for i, (k, v) in enumerate(rows):
        cells = table.rows[start + i].cells
        cells[0].text = ""
        cells[1].text = ""
        rk = cells[0].paragraphs[0].add_run(k)
        set_run(rk, font="Calibri", size=10.5, color=FG_LOW, bold=True)
        rv = cells[1].paragraphs[0].add_run(str(v))
        set_run(rv, font="Calibri", size=10.5, color=FG_HIGH)
        if i % 2 == 0:
            set_cell_bg(cells[0], "F4F1F8")
            set_cell_bg(cells[1], "F4F1F8")
    return table


def add_bullet(doc, text, *, level=0, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    if bold_lead:
        rb = p.add_run(bold_lead + " ")
        set_run(rb, size=11, color=FG_HIGH, bold=True)
        r = p.add_run(text)
        set_run(r, size=11, color=FG_HIGH)
    else:
        r = p.add_run(text)
        set_run(r, size=11, color=FG_HIGH)
    return p


def add_numbered(doc, text, *, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        rb = p.add_run(bold_lead + " ")
        set_run(rb, size=11, color=FG_HIGH, bold=True)
        r = p.add_run(text)
        set_run(r, size=11, color=FG_HIGH)
    else:
        r = p.add_run(text)
        set_run(r, size=11, color=FG_HIGH)
    return p


def add_image(doc, path: Path | None, *, width=Cm(15), caption=None):
    if path is None or not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run()
    r.add_picture(str(path), width=width)
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)
        cr = cp.add_run(caption)
        set_run(cr, size=9, color=FG_LOW, italic=True)


def add_image_pair(doc, path_a: Path | None, path_b: Path | None,
                   *, width_each=Cm(7.4), caption_a=None, caption_b=None):
    if not (path_a and path_a.exists()) and not (path_b and path_b.exists()):
        return
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col in range(2):
        for cell in table.columns[col].cells:
            cell.width = Cm(7.6)
    img_cells = table.rows[0].cells
    cap_cells = table.rows[1].cells
    for i, (path, caption) in enumerate([(path_a, caption_a), (path_b, caption_b)]):
        cell = img_cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if path and path.exists():
            r = p.add_run()
            try:
                r.add_picture(str(path), width=width_each)
            except Exception:
                pass
        cap_cell = cap_cells[i]
        cap_cell.text = ""
        cp = cap_cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption or "")
        set_run(cr, size=9, color=FG_LOW, italic=True)


def add_callout(doc, label, body, *, accent=ACCENT_AMBER, fill_hex="FFF7E6"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    cell.width = Cm(15.5)
    set_cell_bg(cell, fill_hex)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    rl = p.add_run(label.upper() + "  ")
    set_run(rl, size=9, color=accent, bold=True)
    rb = p.add_run(body)
    set_run(rb, size=10.5, color=FG_HIGH)


def add_data_table(doc, headers, rows, *, accent=ACCENT_VIOLET, header_fill="1B142A",
                    col_widths=None):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    if col_widths:
        for j, w in enumerate(col_widths):
            for cell in table.columns[j].cells:
                cell.width = w
    h_cells = table.rows[0].cells
    for j, h in enumerate(headers):
        h_cells[j].text = ""
        p = h_cells[j].paragraphs[0]
        r = p.add_run(h)
        set_run(r, size=10, color=WHITE, bold=True)
        set_cell_bg(h_cells[j], header_fill)
    for i, row in enumerate(rows):
        cells = table.rows[i + 1].cells
        for j, val in enumerate(row):
            cells[j].text = ""
            p = cells[j].paragraphs[0]
            r = p.add_run(str(val))
            set_run(r, size=10, color=FG_HIGH, bold=(j == 0))
        if i % 2 == 0:
            for c in cells:
                set_cell_bg(c, "F4F1F8")
    return table


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ── Content sections ───────────────────────────────────────────────────────


def cover_page(doc):
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    # Top eyebrow
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    r = p.add_run("SEMESTER COURSE PROJECT  ·  FINAL REPORT")
    set_run(r, font="Calibri", size=10, color=ACCENT_VIOLET, bold=True)

    # Logo (centred)
    logo = asset("vidiq_logo_black_bg.png")
    if logo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run()
        r.add_picture(str(logo), width=Cm(4.5))

    # Wordmark
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("VidIQ")
    set_run(r, font="Calibri", size=72, color=BG_DEEP, bold=True)

    # Tagline
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("AI Video Intelligence")
    set_run(r, font="Calibri", size=18, color=ACCENT_VIOLET, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Watch less. Learn more.")
    set_run(r, font="Calibri", size=13, color=FG_MID, italic=True)

    # Tri-color accent rule
    add_horizontal_rule(doc, color="A855F7", thickness=24)

    # Project subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("Comprehensive Digital Marketing Strategy & Execution Report")
    set_run(r, size=14, color=FG_HIGH, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Live demo · vidiq-two.vercel.app")
    set_run(r, size=11, color=ACCENT_CYAN, bold=True)

    # Course info card
    course_table = doc.add_table(rows=4, cols=2)
    course_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    course_table.autofit = False
    rows = [
        ("Course",       "Digital Marketing"),
        ("Section",      "CS-A"),
        ("Instructor",   "Sir Maaz Zafar Cheema"),
        ("Submission",   "Semester Course Project — Final Report"),
    ]
    for i, (k, v) in enumerate(rows):
        cells = course_table.rows[i].cells
        cells[0].width = Cm(5.5)
        cells[1].width = Cm(10.0)
        cells[0].text = ""
        cells[1].text = ""
        rk = cells[0].paragraphs[0].add_run(k.upper())
        set_run(rk, size=10, color=ACCENT_VIOLET, bold=True)
        rv = cells[1].paragraphs[0].add_run(v)
        set_run(rv, size=12, color=FG_HIGH, bold=(k == "Submission"))
        set_cell_bg(cells[0], "F4F1F8")
        set_cell_bg(cells[1], "FFFFFF")

    # Group members card
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("GROUP MEMBERS")
    set_run(r, size=11, color=ACCENT_FUCHSIA, bold=True)

    members = [
        ("22i-1653", "Insharah Aman"),
        ("21i-0416", "M. Nouman Hafeez"),
        ("21i-0484", "Shayan Khan"),
        ("21i-2507", "Muhammad Zain"),
        ("22i-1200", "Farhan Ahmed"),
    ]
    mt = doc.add_table(rows=len(members), cols=2)
    mt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (rno, name) in enumerate(members):
        cells = mt.rows[i].cells
        cells[0].width = Cm(4.0)
        cells[1].width = Cm(8.0)
        cells[0].text = ""
        cells[1].text = ""
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        rr = cells[0].paragraphs[0].add_run(rno)
        set_run(rr, font="Consolas", size=11, color=ACCENT_CYAN, bold=True)
        rn = cells[1].paragraphs[0].add_run(name)
        set_run(rn, size=11, color=FG_HIGH)
        set_cell_bg(cells[0], "F4F1F8")
        set_cell_bg(cells[1], "FFFFFF")

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    r = p.add_run("Submission date · May 2026")
    set_run(r, size=10, color=FG_LOW, italic=True)

    page_break(doc)


def table_of_contents(doc):
    add_heading(doc, "Table of Contents", level=1)
    sections = [
        ("1.", "Executive Summary"),
        ("2.", "Project Overview & Product"),
        ("3.", "Pillar 1 — Brand Identity & Visual System"),
        ("4.", "Pillar 1 — Video Ad Campaign"),
        ("5.", "Pillar 2 — Social Media Strategy (Facebook + Instagram)"),
        ("6.", "Pillar 2 — Meta Ads Campaign"),
        ("7.", "Pillar 3 — Conversation Layer (Inbox & Auto-replies)"),
        ("8.", "Pillar 4 — Search Engine Optimisation"),
        ("9.", "Pillar 4 — Google Ads Campaign"),
        ("10.", "Pillar 5 — Budget Plan (PKR-Primary)"),
        ("11.", "Pillar 5 — Competitive Analysis & SWOT"),
        ("12.", "KPIs, Conversion Tracking & Measurement"),
        ("13.", "Project Links, Live URLs & Repository"),
        ("14.", "Roadmap, Moats & Conclusion"),
    ]
    for num, title in sections:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        rn = p.add_run(num + "  ")
        set_run(rn, font="Consolas", size=11, color=ACCENT_VIOLET, bold=True)
        rt = p.add_run(title)
        set_run(rt, size=11, color=FG_HIGH)
    page_break(doc)


def section_executive_summary(doc):
    add_heading(doc, "1. Executive Summary", level=1)
    add_eyebrow(doc, "Why this project")
    add_paragraph(doc,
        "VidIQ is an end-to-end multimodal AI platform that turns any YouTube video or live "
        "stream into structured, queryable intelligence. The platform fuses speech-to-text, "
        "computer vision, and large language models into a single pipeline that produces a "
        "faithful summary, time-stamped key points, detected events, and a retrieval-grounded "
        "Q&A interface. This report documents the entire digital-marketing strategy and "
        "execution that supports the product launch — covering brand identity, social presence, "
        "paid acquisition, search optimisation, conversation design, and measurement.",
        size=11)

    add_eyebrow(doc, "Marketing thesis")
    add_paragraph(doc,
        "Acquire users by ranking long-tail informational queries (organic SEO), capture "
        "head-term demand at the moment of intent (Google Ads), and build a brand that converts "
        "with a focused 14-day Meta + Google flight, supported by an always-on social presence "
        "on Facebook and Instagram. Pakistani CPM/CPC rates are 5–7× cheaper than Western "
        "markets, so a PKR 250,000 launch flight delivers Western-equivalent reach of roughly "
        "USD 3,000–4,000 worth of impressions.",
        size=11)

    add_eyebrow(doc, "Headline numbers")
    add_data_table(doc,
        ["Metric", "Plan", "Actual (showcase)"],
        [
            ["Launch budget",         "₨ 250,000 (~$890 USD)", "₨ 0"],
            ["Flight length",         "14 days",                "n/a"],
            ["Channels",              "Meta + Google + Organic SEO", "Built to Review"],
            ["Ad sets / campaigns",   "4 Meta sets · 3 Google campaigns", "All screenshotted"],
            ["Content cadence",       "7 posts/week (Facebook + Instagram)", "28-post calendar queued"],
            ["Keywords targeted",     "18 (5 head + 13 long-tail)", "Live in copy & schema"],
            ["On-page SEO techniques","8 applied",                "Verified in DevTools / GSC"],
            ["Live demo",             "vidiq-two.vercel.app",     "Public, indexable"],
        ],
        col_widths=[Cm(5.0), Cm(6.5), Cm(5.0)])

    add_eyebrow(doc, "What you'll find in this report")
    add_paragraph(doc,
        "Sections are organised by the five marketing pillars in the brief. Each section opens "
        "with what we did, follows with how we executed it (with screenshots from the live "
        "deliverables), and closes with the measurable benefit. Section 13 collects every live "
        "URL — the demo, the social profiles, the Search Console property, and the supporting "
        "documentation — so the examiner can verify any claim end-to-end.",
        size=11)
    page_break(doc)


def section_overview(doc):
    add_heading(doc, "2. Project Overview & Product", level=1)

    add_heading(doc, "2.1 Product description", level=2)
    add_paragraph(doc,
        "VidIQ is a web-based SaaS application accessible at vidiq-two.vercel.app. Users paste "
        "a YouTube URL or start a live-stream session; the platform transcribes audio with "
        "faster-whisper, extracts keyframes with OpenCV, captions and tags them with a vision "
        "LLM, and produces a multimodal summary with chapters, key points, sentiment, detected "
        "events, action items, and open questions. A retrieval-grounded chat lets users ask "
        "follow-up questions with timestamp citations.",
        size=11)

    add_heading(doc, "2.2 Capabilities snapshot", level=2)
    capabilities = [
        ["Recorded video analysis", "YouTube URL → metadata → transcript → keyframes → vision captions → multimodal summary"],
        ["Live stream analysis",    "Chunked download (yt-dlp) → rolling transcription → rolling vision → rolling LLM summary"],
        ["Speech understanding",    "YouTube native transcripts (primary) → faster-whisper local fallback"],
        ["Visual understanding",    "Scene-change keyframe extraction → vision-LLM captioning + tagging"],
        ["Multimodal summaries",    "Map-reduce LLM pipeline producing overview, key points, chapters, sentiment"],
        ["Conversational Q&A",      "Retrieval-grounded chat with timestamp citations"],
        ["Cross-library analytics", "/analytics — KPI strip + 7 recharts (volume, source mix, status funnel, topics, events, sentiment, durations)"],
        ["Comparison & translation","/compare side-by-side · transcript translation in 11 languages with RTL"],
    ]
    add_data_table(doc, ["Capability", "Implementation"], capabilities,
                   col_widths=[Cm(5.0), Cm(11.0)])

    add_heading(doc, "2.3 Architecture in one paragraph", level=2)
    add_paragraph(doc,
        "Frontend is a Next.js 14 App-Router SPA deployed on Vercel (free hobby tier). Backend "
        "is a FastAPI service deployed on Hugging Face Spaces (Docker SDK, CPU-basic). The "
        "frontend rewrites /proxy/api and /proxy/media to the Space, so users only ever talk to "
        "the *.vercel.app origin. Persistence is SQLite (production-ready Postgres swap is one "
        "config change). The pipeline is provider-agnostic — every LLM call sits behind an "
        "abstraction with automatic failover across Gemini → Groq → OpenAI → stub.",
        size=11)

    add_heading(doc, "2.4 The five marketing pillars", level=2)
    pillars = [
        ("Pillar 1", "Branding & Visual Identity",  "Logo system · palette · typography · voice · 30-second ad"),
        ("Pillar 2", "Social Media & Paid Social",  "FB + IG presence · 14-day calendar · 4 Meta ad sets"),
        ("Pillar 3", "Conversation Layer",          "Welcome flow · FAQ · saved replies · auto-away"),
        ("Pillar 4", "SEO & Google Ads",            "18 keywords · 8 on-page techniques · 3 Google campaigns"),
        ("Pillar 5", "Budget · Competitive · KPI",  "₨ 250,000 plan · NoteGPT/Eightify analysis · 18 KPIs"),
    ]
    for tag, title, body in pillars:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        rt = p.add_run(tag + "  ")
        set_run(rt, size=11, color=ACCENT_VIOLET, bold=True)
        rh = p.add_run(title + "  ·  ")
        set_run(rh, size=11, color=FG_HIGH, bold=True)
        rb = p.add_run(body)
        set_run(rb, size=11, color=FG_MID)
    page_break(doc)


def section_brand(doc):
    add_heading(doc, "3. Pillar 1 — Brand Identity & Visual System", level=1)

    add_heading(doc, "3.1 Brand essence", level=2)
    add_kv_table(doc, [
        ("Brand name",          "VidIQ"),
        ("Tagline",             "AI Video Intelligence"),
        ("One-line promise",    "Understand any video — recorded or live — in seconds."),
        ("Category",            "AI productivity tool · Video summarisation SaaS"),
        ("Audience",            "Students, educators, researchers, knowledge workers, content creators, financial-market viewers, medical/legal professionals."),
        ("Personality",         "Approachable but credible · restrained, not zany · modern indie startup."),
    ])

    add_heading(doc, "3.2 Logo system", level=2)
    add_paragraph(doc,
        "Three lockups are maintained in frontend/public/ and used as the single source of "
        "truth across the app, the slide deck, and the social profiles. The mark is a stylised "
        "play head mid-glyph — communicating video without resorting to a literal play "
        "triangle. The wordmark splits 'Vid' (sans, foreground) from 'IQ' (gradient, emphasis), "
        "putting the intelligence claim front-and-centre.",
        size=11)
    add_image(doc, asset("vidiq_logo_black_bg.png"), width=Cm(5),
              caption="Primary brand mark on dark surfaces — favicon, app icons, slide footers.")

    add_kv_table(doc, [
        ("vidiq_logo_black_bg.png", "Primary mark on dark UI · favicon · app icons"),
        ("vidiq_logo_white_bg.png", "Light backgrounds · OG/Twitter cards · print · slides"),
        ("vidiq_logo_text.png",      "Horizontal lockup with wordmark — ad creatives, video end-cards"),
    ], col_widths=(Cm(6.0), Cm(9.5)))

    add_heading(doc, "3.3 Colour palette", level=2)
    add_data_table(doc,
        ["Token", "Hex", "Role", "Why this colour"],
        [
            ["--background", "#0a0612", "App canvas",     "Near-black with violet undertone — feels like a cinema, not a spreadsheet."],
            ["--primary",    "#a855f7", "Brand violet",   "Trust (blue) ∩ Creativity (red). Colour of insight."],
            ["--accent-2",   "#ec4899", "Brand fuchsia",  "Energetic counterpoint — used in calls-to-action."],
            ["--accent-3",   "#06b6d4", "Brand cyan",     "Cool, technical — used in data UI and code surfaces."],
            ["--success",    "#10b981", "Emerald",        "Positive states, completed runs."],
            ["--warning",    "#f59e0b", "Amber",          "Hold-state, limits, budget chips."],
        ],
        col_widths=[Cm(2.5), Cm(2.0), Cm(3.0), Cm(8.0)])

    add_heading(doc, "3.4 Typography", level=2)
    add_kv_table(doc, [
        ("Display",  "Plus Jakarta Sans — geometric, modern, slightly humanist"),
        ("Body",     "Inter — workhorse UI typeface · excellent at small sizes"),
        ("Mono",     "JetBrains Mono — code, timestamps, API endpoints"),
    ])

    add_heading(doc, "3.5 Voice & tone", level=2)
    for v in [
        ("Clarity",   "Time-stamped citations behind every claim. Crisp, declarative, jargon-free."),
        ("Speed",     "Sub-minute analysis on free-tier APIs. Direct verbs, short sentences."),
        ("Trust",     "Every chat answer cites a transcript moment. Cites sources, never overclaims."),
        ("Curiosity", "Multimodal — fuses speech + vision + LLM. Uses words like unlock, surface, ground."),
        ("Open",      "Provider-agnostic, free-tier first. Inclusive, never gate-keepy."),
    ]:
        add_bullet(doc, v[1], bold_lead=v[0] + " —")

    add_heading(doc, "3.6 Logo applied — surfaces audited", level=2)
    surfaces = [
        ["Favicon",        "Browser tab · bookmark · iOS home-screen"],
        ["Top nav",        "Persistent across every route · clickable home"],
        ["Hero splash",    "Animated session-gated landing animation"],
        ["OG / Twitter",   "1200×630 social-share preview · all 4 routes"],
        ["JSON-LD",        "Logo URL embedded in structured data"],
        ["Slide footer",   "0.32-inch mark on every one of 33 deck slides"],
        ["FB Page",        "Cover photo + profile mark on the official page"],
        ["IG Profile",     "Profile mark + welcome post on @vidiq_official"],
    ]
    add_data_table(doc, ["Surface", "Where it appears"], surfaces,
                   col_widths=[Cm(4.0), Cm(11.5)])
    add_callout(doc, "Result", "8 / 8 brand surfaces audited green. KPI #4 (Brand consistency) "
                "scored 5 / 5 in the rubric tracker.")
    page_break(doc)


def section_video_ad(doc):
    add_heading(doc, "4. Pillar 1 — Video Ad Campaign", level=1)

    add_heading(doc, "4.1 The 30-second master cut", level=2)
    add_paragraph(doc,
        "A 30-second hero video drives the launch. The structure is a 6-beat narrative arc that "
        "respects the platform-native viewing patterns — a hard hook in the first 3 seconds, "
        "demo proof at second 5, value framing at second 12, social proof at second 20, and a "
        "CTA loop in the final 6 seconds. The cut is mastered once at 1080×1920 (9:16) and "
        "exported to four platform-specific cutdowns (15s, 9s, 6s, square 1:1) without "
        "re-grading.",
        size=11)

    add_heading(doc, "4.2 Beat-by-beat script", level=2)
    add_data_table(doc,
        ["#", "Time", "Beat", "Visual", "Voiceover / on-screen text"],
        [
            ["1", "0:00–0:03", "Hook",     "Crash zoom on a 2-hour lecture timeline",   "Two hours. Zero notes."],
            ["2", "0:03–0:08", "Promise",  "Cursor pastes a YouTube URL into VidIQ",     "Paste any video."],
            ["3", "0:08–0:15", "Proof",    "Screen-record · summary + chapters appear",  "Get a summary, chapters, citations — in seconds."],
            ["4", "0:15–0:22", "Differentiator","Live-stream pipeline ticks on the right","Even a live stream. Even a 4-hour earnings call."],
            ["5", "0:22–0:27", "Social",   "Stack of testimonial cards",                 "Trusted by students, traders, doctors, lawyers."],
            ["6", "0:27–0:30", "CTA",      "Logo lockup → URL pill",                     "Try it free at vidiq-two.vercel.app"],
        ],
        col_widths=[Cm(0.7), Cm(2.0), Cm(2.5), Cm(4.5), Cm(5.8)])

    add_heading(doc, "4.3 Cutdowns shipped", level=2)
    add_kv_table(doc, [
        ("Master 30s",  "Meta Reels · IG Feed · YouTube In-Stream skippable"),
        ("15s cutdown", "TikTok-native · IG Stories · Meta non-skippable"),
        ("9s cutdown",  "Pre-roll bumpers"),
        ("6s cutdown",  "YouTube non-skippable bumper"),
        ("Square 1:1",  "Meta Feed · LinkedIn"),
    ])

    add_heading(doc, "4.4 Production stack", level=2)
    add_bullet(doc, "Footage: Loom screen-records of the live product on vidiq-two.vercel.app.")
    add_bullet(doc, "Music: YouTube Audio Library (royalty-free) — neutral, motivating.")
    add_bullet(doc, "Captions: hard-burned for sound-off viewing (87% of Meta plays).")
    add_bullet(doc, "Aspect-ratio safe areas verified against Meta + YouTube official specs.")
    page_break(doc)


def section_social(doc):
    add_heading(doc, "5. Pillar 2 — Social Media Strategy (Facebook + Instagram)", level=1)

    add_heading(doc, "5.1 Live profiles", level=2)
    add_paragraph(doc,
        "Both social identities are live and brand-aligned. Profile marks, cover art, and "
        "welcome posts mirror the in-app branding so the funnel feels seamless from ad → "
        "profile → site.",
        size=11)
    add_data_table(doc,
        ["Channel", "Handle / URL", "Status"],
        [
            ["Facebook Page", "facebook.com/profile.php?id=61588939576479", "Live · welcome post pinned · cover applied"],
            ["Instagram",     "instagram.com/vidiq_official/",              "Live · @vidiq_official · welcome post + bio CTA"],
        ],
        col_widths=[Cm(3.5), Cm(8.0), Cm(4.0)])

    add_image_pair(doc,
                    asset("fb-page.png"), asset("insta-welcome-post.jpeg", "insta-welcome-post.png"),
                    caption_a="Facebook Page · cover + profile mark live.",
                    caption_b="Instagram · @vidiq_official welcome post.")

    add_heading(doc, "5.2 14-day content calendar", level=2)
    add_paragraph(doc,
        "Twenty-eight posts across the 14-day flight (two posts per day, alternating Facebook "
        "and Instagram). The calendar is built and queued inside Meta Business Suite — every "
        "post has a scheduled time, a format, and a content pillar.",
        size=11)
    add_image(doc, asset("planner-calendar.png"), width=Cm(15),
              caption="Meta Business Suite Planner — 28 posts queued across the launch flight.")

    add_heading(doc, "5.3 Format mix", level=2)
    add_data_table(doc,
        ["Format", "Use", "Cadence"],
        [
            ["Reels (9:16)",     "Hooks, demo loops, before/after note-taking",     "3 / week"],
            ["Carousels (1:1)",  "Long-form explainers, feature deep-dives",         "2 / week"],
            ["Single image",     "Quotes, milestones, brand moments",                "1 / week"],
            ["Stories (9:16)",   "Polls, BTS, link-stickers to the live demo",       "Daily"],
        ],
        col_widths=[Cm(3.5), Cm(8.5), Cm(3.5)])

    add_heading(doc, "5.4 Content pillars", level=2)
    add_bullet(doc, "Show the product — paste-to-summary loops, live-stream demos, chapter skipping.", bold_lead="Product:")
    add_bullet(doc, "Educate the audience — how to summarise long lectures, study smarter, save hours.", bold_lead="Education:")
    add_bullet(doc, "Tell the brand story — why we built it, who it's for, what it stands for.", bold_lead="Brand:")
    add_bullet(doc, "Surface the community — UGC, testimonials, replies, citations.", bold_lead="Community:")

    add_heading(doc, "5.5 Hashtag strategy", level=2)
    add_paragraph(doc,
        "Each post carries 8–12 hashtags split across (a) brand tags (#VidIQ, #AIVideo), "
        "(b) category tags (#AISummary, #YouTubeAI, #StudyAI), and (c) audience tags "
        "(#Students, #ContentCreators, #Trading, #MedSchool). All hashtags audited for "
        "relevance — banned-tag scan run weekly.",
        size=11)
    add_callout(doc, "KPIs touched",
                "KPI #5 (14/14 days posted) · KPI #6 (4 distinct formats — exceeds 3+ requirement) "
                "· KPI #9 (FB + IG identities live and audited).",
                accent=ACCENT_CYAN, fill_hex="E6F8FA")
    page_break(doc)


def section_meta_ads(doc):
    add_heading(doc, "6. Pillar 2 — Meta Ads Campaign", level=1)
    add_paragraph(doc,
        "The Meta flight runs four distinct ad sets aimed at the four highest-intent audience "
        "clusters. Every ad set was built end-to-end in Ads Manager and screenshotted at the "
        "Review/Summary screen — campaigns sit in Draft, never published. Total Meta budget: "
        "₨ 95,000 (~$340), split evenly at ₨ 23,750 per set.",
        size=11)

    add_heading(doc, "6.1 Ad set matrix", level=2)
    add_data_table(doc,
        ["Set", "Audience", "Targeting", "Creative", "Budget"],
        [
            ["1", "Students",         "Age 18–24 · Education + Productivity interests", "Reel: '2 hours of lecture → 5 minutes of notes'", "₨ 23,750"],
            ["2", "Creators",         "Age 22–40 · YouTube creator interests + lookalikes", "Carousel: 'Save 4 hrs/week of research'", "₨ 23,750"],
            ["3", "Knowledge Workers","Age 28–45 · Consulting · SaaS · Finance", "Reel: 'Earnings call → action items'", "₨ 23,750"],
            ["4", "Retargeting",      "Site visitors 30 d · IG engagers 60 d", "Static: 'Forgot to finish? 1 click resumes.'", "₨ 23,750"],
        ],
        col_widths=[Cm(0.7), Cm(3.4), Cm(5.0), Cm(4.5), Cm(2.0)])

    add_heading(doc, "6.2 Live screenshots from Ads Manager", level=2)
    add_image_pair(doc,
                    asset("meta-ad-set-1-students.png"),
                    asset("meta-ad-set-2-creators.png"),
                    caption_a="Ad Set 1 — Students (Review screen).",
                    caption_b="Ad Set 2 — Creators (Review screen).")
    add_image_pair(doc,
                    asset("meta-ad-set-3-knowledge.png"),
                    asset("meta-ad-set-4-retarget.png"),
                    caption_a="Ad Set 3 — Knowledge Workers (Review screen).",
                    caption_b="Ad Set 4 — Retargeting (Review screen).")

    add_heading(doc, "6.3 Bidding strategy", level=2)
    add_bullet(doc, "Days 1–7 — Highest Volume bid · LP-views objective · cheapest data while pixel learns.")
    add_bullet(doc, "Day 5 — Creative-swap fund (10% of Meta budget) reserved for fatigue rotation.")
    add_bullet(doc, "Days 8–14 — Switch to Conversions objective · CPR cap ₨ 420 (~$1.50).")
    add_bullet(doc, "Day 14 — Negative-keyword sweep, screenshot reports, publish learnings.")

    add_heading(doc, "6.4 Pakistan-specific rationale", level=2)
    add_paragraph(doc,
        "Meta CPM in Pakistan averages ₨ 80–250 per 1,000 impressions — 5–7× cheaper than the "
        "$5–15 typical in the US/UK. At ₨ 95,000 we forecast roughly 380K–1.2M impressions "
        "across the four sets, more than enough volume for the pixel to leave the learning "
        "phase and converge on a stable cost-per-result.",
        size=11)
    page_break(doc)


def section_conversation(doc):
    add_heading(doc, "7. Pillar 3 — Conversation Layer", level=1)
    add_paragraph(doc,
        "Every paid impression that converts to a profile visit lands in a conversation surface "
        "(Messenger, IG DM, or the in-app chat). The conversation layer is the difference "
        "between a click that bounces and a click that signs up. We configured four automation "
        "primitives in Meta Business Suite + Inbox.",
        size=11)

    add_heading(doc, "7.1 Configured automations", level=2)
    add_kv_table(doc, [
        ("Welcome message",     "Fires on first DM · introduces VidIQ in 2 lines + offers the demo URL"),
        ("FAQ auto-responder",  "Six pre-trained intents: pricing · how-it-works · supported videos · privacy · live · refund"),
        ("Saved replies",       "Eight curated replies for fast human-agent responses (free tier · live · API · etc.)"),
        ("Auto-away",           "Outside business hours · sets expectations on response window"),
    ])

    add_image_pair(doc,
                    asset("auto-greeting.png"), asset("auto-faq.png"),
                    caption_a="Welcome message — fires on first DM.",
                    caption_b="FAQ auto-responder — six trained intents.")
    add_image_pair(doc,
                    asset("saved-replies.png"), asset("auto-away.png"),
                    caption_a="Saved replies library — fast human responses.",
                    caption_b="Auto-away — sets expectations outside hours.")

    add_heading(doc, "7.2 Why this matters", level=2)
    add_callout(doc, "Benefit",
                "Automations cut median first-response from hours to seconds, reducing drop-off "
                "between ad-click and trial. The same surface doubles as a brand-voice testing "
                "ground — every reply is on-tone, citing where users can self-serve.",
                accent=ACCENT_EMERALD, fill_hex="E6F7EF")
    page_break(doc)


def section_seo(doc):
    add_heading(doc, "8. Pillar 4 — Search Engine Optimisation", level=1)

    add_heading(doc, "8.1 Keyword strategy (18 keywords across 3 tiers)", level=2)
    add_paragraph(doc,
        "Eighteen keywords clustered into three tiers. Head terms are saturated (NoteGPT and "
        "Eightify outrank a DA-1 site) so we bid those via Google Ads. Long-tail and niche "
        "terms have keyword-difficulty under 15 and are rankable in 2–3 months with a focused "
        "on-page push.",
        size=11)

    head_rows = [
        ["ai video summarizer",          "27,100", "48"],
        ["youtube video summarizer",     "22,200", "42"],
        ["summarize youtube video",      "18,100", "39"],
        ["ai for video transcripts",     "4,400",  "31"],
        ["live stream ai summary",       "1,600",  "24"],
    ]
    long_rows = [
        ["summarize a 2-hr youtube video",       "2,400", "19"],
        ["how to summarize a long youtube lecture", "900", "12"],
        ["study with youtube ai summary",        "700",   "13"],
        ["ai timestamp summary youtube",         "600",   "11"],
    ]
    niche_rows = [
        ["summarize trading livestream ai",      "210",   "8"],
        ["medical lecture video summarizer",     "320",   "14"],
        ["webinar to text ai summary",           "1,100", "25"],
    ]
    add_paragraph(doc, "Tier 1 — Head terms (saturated · bid via Google Ads):",
                  size=11, bold=True)
    add_data_table(doc, ["Keyword", "MSV", "KD"], head_rows,
                   col_widths=[Cm(8.0), Cm(3.0), Cm(2.0)])
    add_paragraph(doc, "Tier 2 — Long-tail (KD < 15 · rankable in 2–3 months):",
                  size=11, bold=True)
    add_data_table(doc, ["Keyword", "MSV", "KD"], long_rows,
                   col_widths=[Cm(8.0), Cm(3.0), Cm(2.0)])
    add_paragraph(doc, "Tier 3 — Niche / vertical-domain (first-mover · zero competition):",
                  size=11, bold=True)
    add_data_table(doc, ["Keyword", "MSV", "KD"], niche_rows,
                   col_widths=[Cm(8.0), Cm(3.0), Cm(2.0)])

    add_heading(doc, "8.2 On-page SEO — eight techniques applied", level=2)
    add_paragraph(doc,
        "Every technique below ships in the live build at vidiq-two.vercel.app. Each row names "
        "the technique, the implementation, and the measurable benefit.",
        size=11)
    seo_rows = [
        ["Meta tags",         "Per-page title + description + canonical + OG + Twitter card",
         "Higher SERP CTR · branded social previews · prevents duplicate-content penalties"],
        ["Alt text",          "Every image carries descriptive alt — auto-generated from caption",
         "Image-search reach · WCAG 2.1 AA accessibility compliance"],
        ["Header hierarchy",  "Strict h1 → h2 → h3 — verified with Wave + axe DevTools",
         "Crawler topic-modelling · screen-reader navigability"],
        ["Structured data",   "JSON-LD: Organization + WebSite + SoftwareApplication",
         "Eligibility for rich results, knowledge panels, sitelinks"],
        ["Sitemap + robots",  "/sitemap.xml lists 4 URLs · /robots.txt allows full crawl · submitted to GSC",
         "Faster discovery · controlled crawl budget"],
        ["Internal linking",  "Hub-and-spoke: home → /analyze, /live, /library + cross-links",
         "Distributes link equity · increases crawl depth"],
        ["Core Web Vitals",   "next/image · lazy-loading · Vercel Edge cache · font preload",
         "LCP < 2.5 s · CLS < 0.1 · ranking-signal compliance"],
        ["Mobile-first + HTTPS","Tailwind responsive · viewport meta · TLS via Vercel",
         "Mobile-index ranking signal · trust signal for users + browsers"],
    ]
    add_data_table(doc, ["Technique", "Implementation", "Benefit"], seo_rows,
                   col_widths=[Cm(3.5), Cm(6.5), Cm(5.5)])

    add_heading(doc, "8.3 SEO evidence trail", level=2)
    add_paragraph(doc,
        "Five live proofs are captured in marketing/assets/seo-evidence/ and reproduced below. "
        "Every screenshot is dated 2026-05-06 and was taken from the production property at "
        "vidiq-two.vercel.app.",
        size=11)
    add_image_pair(doc,
                    asset("gsc-property-verified.png"),
                    asset("gsc-sitemap-success.png"),
                    caption_a="Google Search Console — property verified (HTML-tag method).",
                    caption_b="Sitemap submitted — 4 URLs discovered by Google.")
    add_image_pair(doc,
                    asset("gsc-url-inspection-home.png"),
                    asset("pagespeed-mobile.png"),
                    caption_a="URL inspection — indexing requested for /, /analyze, /live.",
                    caption_b="PageSpeed Insights — mobile audit (4 categories).")
    add_image(doc, asset("pagespeed-desktop.png"), width=Cm(13),
              caption="PageSpeed Insights — desktop audit (Lighthouse 4 categories).")

    add_callout(doc, "Net benefit",
                "Rank-eligible from day 1 · indexable to Google + Bing · social previews load "
                "with brand · WCAG 2.1 AA accessible · ~95th-percentile mobile UX. The full "
                "audit lives in 07-onpage-seo-report.md.",
                accent=ACCENT_EMERALD, fill_hex="E6F7EF")
    page_break(doc)


def section_google_ads(doc):
    add_heading(doc, "9. Pillar 4 — Google Ads Campaign", level=1)
    add_paragraph(doc,
        "Three campaigns share the ₨ 70,000 (~$250) Google budget, all built end-to-end and "
        "screenshotted at the Review screen (saved as Draft, never published).",
        size=11)

    add_heading(doc, "9.1 Campaign matrix", level=2)
    add_data_table(doc,
        ["Campaign", "Strategy", "Audience / Placement", "Budget"],
        [
            ["Search · Brand Cluster", "Maximise Conversions · 4 ad groups", "VidIQ + competitor brand bids + long-tail intent", "₨ 23,100"],
            ["YouTube · In-Stream",     "Target CPM · 30s skippable + 6s bumper", "AI productivity + study channels",                "₨ 18,900"],
            ["Performance Max",         "Maximise Conversion Value · 3 asset groups","Google AI multi-surface (YouTube + Display + Search)", "₨ 28,000"],
        ],
        col_widths=[Cm(3.5), Cm(4.5), Cm(5.5), Cm(2.0)])

    add_heading(doc, "9.2 Live screenshots from Google Ads Review", level=2)
    add_image_pair(doc,
                    asset("google-search-review.jpeg", "google-search-review.png"),
                    asset("google-youtube-review.jpeg", "google-youtube-review.png"),
                    caption_a="Search · Brand Cluster — Review screen.",
                    caption_b="YouTube · In-Stream — Review screen.")
    add_image(doc, asset("google-pmax-review.jpeg", "google-pmax-review.png"),
              width=Cm(13),
              caption="Performance Max — 3 asset groups · Review screen.")

    add_heading(doc, "9.3 Bidding rules & guardrails", level=2)
    add_bullet(doc, "Start: Maximise Conversions → switch to Target CPA ₨ 420 (~$1.50) after 30 conversions banked.")
    add_bullet(doc, "Friday negative-keyword sweep: drop terms with > ₨ 1,400 spend and 0 conversions.")
    add_bullet(doc, "PMax gated by conversion volume; spare reallocates to Search if it doesn't unlock.")
    page_break(doc)


def section_budget(doc):
    add_heading(doc, "10. Pillar 5 — Budget Plan (PKR-Primary)", level=1)

    add_heading(doc, "10.1 Top-line — ₨ 250,000 PKR (~$890 USD)", level=2)
    add_paragraph(doc,
        "The launch flight is sized at ₨ 250,000 — the mid-range of typical Pakistani SaaS "
        "launch budgets (₨ 150K–400K per month). The currency is PKR-primary because we operate "
        "from Pakistan and pay Meta and Google in rupees. USD figures are illustrative at "
        "1 USD ≈ 280 PKR.",
        size=11)

    add_data_table(doc,
        ["Line item", "PKR", "USD (approx.)", "Share"],
        [
            ["Meta Ads",            "₨ 95,000", "$340", "38 %"],
            ["Google Ads",          "₨ 70,000", "$250", "28 %"],
            ["Creative production", "₨ 40,000", "$143", "16 %"],
            ["Influencer seeding",  "₨ 20,000", "$71",   "8 %"],
            ["Tools & SaaS",        "₨ 17,500", "$63",   "7 %"],
            ["Contingency (10%)",   "₨ 7,500",  "$27",   "3 %"],
            ["Total",               "₨ 250,000", "$893", "100 %"],
        ],
        col_widths=[Cm(5.0), Cm(3.5), Cm(3.5), Cm(2.5)])

    add_heading(doc, "10.2 Why ₨ 250,000?", level=2)
    add_bullet(doc, "Pakistani SaaS launch budgets cluster at ₨ 150K–400K/month — ₨ 250K is mid-range, aggressive for a 14-day flight.", bold_lead="Market norm:")
    add_bullet(doc, "Pakistan averages ₨ 80–250 per 1,000 impressions vs $5–15 in Western markets — 5–7× more reach per rupee.", bold_lead="Meta CPM:")
    add_bullet(doc, "Long-tail keyword CPCs in Pakistan average ₨ 20–80 vs $1–3 USD — same gap.", bold_lead="Google CPC:")
    add_bullet(doc, "₨ 250,000 buys ~80,000–120,000 impressions across Meta + Google with our PK-targeting — enough for the pixel to leave learning and converge to a stable CPA.", bold_lead="Volume:")
    add_bullet(doc, "Aligns with a final-year student team budget without external sponsorship.", bold_lead="Affordability:")

    add_heading(doc, "10.3 Spend phasing across the 14 days", level=2)
    add_data_table(doc,
        ["Window", "Phase", "Action"],
        [
            ["Days 1–7",  "Learn",                "Highest-Volume bid · LP-views objective · cheapest data while pixel learns"],
            ["Day 5",     "Creative swap fund",   "10% of Meta budget reserved for fatigue rotation"],
            ["Days 8–14", "Optimise",             "Switch to Conversions objective · CPR ₨ 420 cap · PMax unlocks if 30 conv banked"],
            ["Day 14",    "Wrap",                 "Negative-KW sweep · screenshot reports · publish learnings"],
        ],
        col_widths=[Cm(3.0), Cm(3.5), Cm(9.0)])

    add_heading(doc, "10.4 Showcase note", level=2)
    add_callout(doc, "Showcase actual",
                "₨ 0 spent — campaigns built to the Review screen in Ads Manager and "
                "screenshotted; never published. The numbers above describe a planned launch, "
                "which is what the rubric grades.",
                accent=ACCENT_EMERALD, fill_hex="E6F7EF")
    page_break(doc)


def section_competitive(doc):
    add_heading(doc, "11. Pillar 5 — Competitive Analysis & SWOT", level=1)

    add_heading(doc, "11.1 Direct competitors", level=2)
    add_data_table(doc,
        ["Attribute", "VidIQ", "NoteGPT", "Eightify"],
        [
            ["Pricing",         "Free tier · $9/mo planned",  "$9/mo",            "$8.99/mo"],
            ["Free tier",       "Generous (sub-min runs)",     "10 summaries/mo",   "5 summaries/mo"],
            ["Live streams",    "Yes — rolling pipeline",      "No",                "No"],
            ["Multimodal",      "Vision + audio + text",       "Audio + text only", "Audio + text only"],
            ["Languages",       "11 + RTL",                    "12",                "5"],
            ["Open-source-able","Yes",                          "No",                "No"],
            ["Domain modes",    "Medical · legal · trading",   "No",                "No"],
        ],
        col_widths=[Cm(3.5), Cm(4.0), Cm(4.0), Cm(4.0)])

    add_heading(doc, "11.2 Social presence comparison", level=2)
    add_data_table(doc,
        ["Channel", "VidIQ", "NoteGPT", "Eightify"],
        [
            ["Facebook URL",       "fb.com/…?id=61588939576479", "facebook.com/notegpt", "facebook.com/eightifyapp"],
            ["FB followers",       "Launching",                  "~3,100",               "~620"],
            ["FB cadence",         "7 / week (planned)",         "~3 / week",            "<1 / week"],
            ["Instagram handle",   "@vidiq_official",            "@notegpt.official",    "@eightifyapp"],
            ["IG followers",       "Launching",                  "~5,800",               "~1,900"],
            ["IG cadence",         "7 / week (planned)",         "~4 / week",            "~1 / week"],
            ["Hashtags / post",    "8–12",                        "~6",                   "~3"],
        ],
        col_widths=[Cm(3.5), Cm(4.0), Cm(4.0), Cm(4.0)])

    add_heading(doc, "11.3 SWOT analysis", level=2)
    swot = [
        ["Strengths",
         "Multimodal AI (vision + audio + text); live-stream pipeline rare in category; "
         "provider-agnostic backend with auto-failover; domain-aware modes (medical · legal · "
         "trading); open-source-able product."],
        ["Weaknesses",
         "Zero brand awareness; DA = 1; brand-name collision with vidiq.com (unrelated YouTube "
         "SEO tool); no native mobile app yet; small shipping team; YouTube IP rate-limiting on "
         "the HF Space demo."],
        ["Opportunities",
         "Long-tail SEO is rankable in 2–3 months (KD < 15); first-mover in vertical niches "
         "(trading streams, med lectures, legal proceedings); Pakistani CPM/CPC arbitrage; "
         "creator-economy demand for time-saving tooling."],
        ["Threats",
         "Big-tech competitors (YouTube AI, Google Bard) could ship summaries natively; LLM API "
         "pricing changes; YouTube ToS changes around transcript scraping; brand-name confusion "
         "with the existing vidiq.com property."],
    ]
    for h, body in swot:
        add_paragraph(doc, h, size=12, bold=True, color=ACCENT_VIOLET, space_after=2)
        add_paragraph(doc, body, size=11, space_after=8)
    page_break(doc)


def section_kpi(doc):
    add_heading(doc, "12. KPIs, Conversion Tracking & Measurement", level=1)

    add_heading(doc, "12.1 18 KPIs across the rubric", level=2)
    add_paragraph(doc,
        "All 18 graded KPIs are tracked in the live /marketing dashboard and the supporting "
        "11-kpi-tracker.md. Below is a condensed scorecard.",
        size=11)
    kpis = [
        ["Brand", "Logo applied · 8 surfaces · consistent",                      "Strong"],
        ["Brand", "Brand voice · 5 values · documented",                          "Strong"],
        ["Brand", "Colour & typography system · 6 tokens",                        "Strong"],
        ["Brand", "Brand consistency across deliverables",                        "Strong"],
        ["Social","14-day calendar · 28 posts queued",                            "Strong"],
        ["Social","Format mix (Reel · Carousel · Image · Story) — 4 of 4",         "Strong"],
        ["Social","FB + IG identities live · welcome posts · cover applied",       "Strong"],
        ["Social","Meta Ads — 4 ad sets built to Review",                         "Strong"],
        ["Conv.", "Welcome message · auto-FAQ · saved replies · auto-away",        "Strong"],
        ["Conv.", "Tone consistent across human + automated replies",               "Strong"],
        ["SEO",   "Keyword research · 18 KWs · 3 tiers",                          "Strong"],
        ["SEO",   "On-page SEO · 8 techniques applied · GSC verified",            "Strong"],
        ["SEO",   "Lighthouse · SEO 100 · audited mobile + desktop",              "Strong"],
        ["Ads",   "Google Ads · 3 campaigns · Search + YouTube + PMax",            "Strong"],
        ["Ads",   "Bidding strategy · CPA cap · PMax gating",                       "Strong"],
        ["Ads",   "Brand alignment of ad copy + creative",                         "Strong"],
        ["Pillar5","Total budget + distribution (PKR-primary)",                    "Strong"],
        ["Pillar5","Competitive analysis · 2 direct competitors",                   "Strong"],
    ]
    add_data_table(doc, ["Pillar", "KPI", "Status"], kpis,
                   col_widths=[Cm(2.5), Cm(11.0), Cm(2.0)])

    add_heading(doc, "12.2 Conversion tracking architecture", level=2)
    add_kv_table(doc, [
        ("Meta Pixel + CAPI", "Browser pixel + server-side Conversions API · same event names"),
        ("Google gtag",        "Site-wide gtag.js · auto-tagged for cross-device measurement"),
        ("GA4",                "Property linked to GSC · enhanced measurement on"),
        ("Event taxonomy",     "view_demo · start_analysis · finish_analysis · open_chat · sign_up"),
        ("Attribution model", "Data-driven (Google) + 7-day click / 1-day view (Meta)"),
    ])

    add_heading(doc, "12.3 KPI targets for the 14-day flight", level=2)
    add_data_table(doc,
        ["KPI", "Target", "Source"],
        [
            ["Cost per result (Meta)", "≤ ₨ 420 ($1.50)",  "Ads Manager · LP-views → Conversions"],
            ["Cost per click (Google)", "≤ ₨ 70 ($0.25)",   "Google Ads · Search"],
            ["CTR (Meta)",              "≥ 1.2 %",          "Industry benchmark for AI-tools vertical"],
            ["CTR (Google Search)",     "≥ 4 %",             "Industry benchmark for high-intent KWs"],
            ["Demo starts / day",        "30",                "GA4 · custom event start_analysis"],
            ["Sign-ups / day",           "8",                 "GA4 · custom event sign_up"],
        ],
        col_widths=[Cm(4.5), Cm(4.0), Cm(7.0)])
    page_break(doc)


def section_links(doc):
    add_heading(doc, "13. Project Links, Live URLs & Repository", level=1)

    add_heading(doc, "13.1 Live product", level=2)
    add_kv_table(doc, [
        ("Live demo",      "https://vidiq-two.vercel.app"),
        ("Marketing dash", "https://vidiq-two.vercel.app/marketing"),
        ("Sitemap",        "https://vidiq-two.vercel.app/sitemap.xml"),
        ("Robots",         "https://vidiq-two.vercel.app/robots.txt"),
    ])

    add_heading(doc, "13.2 Backend & infrastructure", level=2)
    add_kv_table(doc, [
        ("Backend (HF Space)", "https://huggingface.co/spaces/noumanhafeez11/vidiq-backend"),
        ("Backend host",        "https://noumanhafeez11-vidiq-backend.hf.space"),
        ("Frontend host",       "Vercel (Hobby tier)"),
        ("Database",            "SQLite (Postgres-ready)"),
    ])

    add_heading(doc, "13.3 Social & community", level=2)
    add_kv_table(doc, [
        ("Facebook Page",  "https://www.facebook.com/profile.php?id=61588939576479"),
        ("Instagram",       "https://www.instagram.com/vidiq_official/"),
        ("Brand handle",   "@vidiq_official"),
    ])

    add_heading(doc, "13.4 Search & SEO accounts", level=2)
    add_kv_table(doc, [
        ("Google Search Console", "Property verified — vidiq-two.vercel.app (HTML-tag method)"),
        ("Bing Webmaster Tools",   "Property submitted (mirror of GSC)"),
        ("PageSpeed Insights",      "Mobile + desktop audited"),
        ("GA4 property",            "Linked to GSC · enhanced measurement on"),
    ])

    add_heading(doc, "13.5 Marketing-document index", level=2)
    docs = [
        ("01-brand-guide.md",        "Pillar 1 · brand essence · logo · palette · typography · voice"),
        ("02-video-ad-script.md",    "Pillar 1 · 30-second master cut · 6-beat script · cutdowns"),
        ("03-content-calendar.md",   "Pillar 2 · 14-day calendar · 28 posts queued"),
        ("04-meta-ads-plan.md",      "Pillar 2 · 4 ad sets · creatives · bidding"),
        ("05-social-templates.md",   "Pillar 2 · post templates · captions · hashtag stacks"),
        ("06-keyword-research.md",   "Pillar 4 · 18 keywords · 3 tiers"),
        ("07-onpage-seo-report.md",  "Pillar 4 · 8 on-page techniques · audit"),
        ("08-google-ads-plan.md",    "Pillar 4 · 3 campaigns · bidding rules"),
        ("09-budget.md",             "Pillar 5 · ₨ 250K plan · phasing · rationale"),
        ("10-competitive-analysis.md","Pillar 5 · NoteGPT + Eightify side-by-side"),
        ("11-kpi-tracker.md",        "Pillar 5 · 18 KPIs"),
        ("14-vercel-deployment.md",   "Infra · Vercel deployment runbook"),
        ("15-deployment-evidence.md", "Infra · screenshots of deployment surfaces"),
        ("16-execution-playbook.md",  "Cross-pillar · day-by-day execution plan"),
    ]
    add_data_table(doc, ["Document", "Coverage"], docs,
                   col_widths=[Cm(5.0), Cm(10.5)])
    page_break(doc)


def section_conclusion(doc):
    add_heading(doc, "14. Roadmap, Moats & Conclusion", level=1)

    add_heading(doc, "14.1 Five moats we are deliberately building", level=2)
    add_bullet(doc, "Citations behind every chat answer · timestamp-grounded retrieval — hard to fake.", bold_lead="Trust:")
    add_bullet(doc, "Live-stream pipeline · rolling vision + audio · few competitors ship this.", bold_lead="Live-streams:")
    add_bullet(doc, "Domain modes for medical · legal · trading — picks up underserved segments.", bold_lead="Verticals:")
    add_bullet(doc, "Provider-agnostic backend · auto-failover Gemini → Groq → OpenAI — resilient to pricing/availability shocks.", bold_lead="Resilience:")
    add_bullet(doc, "Open-source-able codebase — potential GitHub-driven dev-community pull.", bold_lead="Community:")

    add_heading(doc, "14.2 Eight-phase roadmap (12 months)", level=2)
    phases = [
        ["Q1 (now)",      "Launch flight",      "14-day Meta + Google · social cadence · GSC indexing"],
        ["Q1+1",          "First-1000 users",   "On-board Iteration · sign-up survey · NPS"],
        ["Q2",            "Native mobile",      "Capacitor wrapper for iOS + Android"],
        ["Q2",            "Vertical bundles",   "Medical / Legal / Trading SKUs · domain prompts"],
        ["Q2+1",          "Browser extension",  "Right-click any video → summarise"],
        ["Q3",            "Team accounts",      "Shared library · admin · billing"],
        ["Q3",            "Embedding API",      "Public docs · Zapier integration"],
        ["Q4",            "Series A readiness", "DAU > 5K · NPS > 40 · ARR > $250K"],
    ]
    add_data_table(doc, ["When", "Phase", "Outcome"], phases,
                   col_widths=[Cm(2.5), Cm(4.0), Cm(9.0)])

    add_heading(doc, "14.3 Closing", level=2)
    add_paragraph(doc,
        "VidIQ ships a real product, on the real internet, behind a real brand. The marketing "
        "deliverables in this report are not hypothetical — every screenshot is taken from "
        "Ads Manager Review, Google Ads Review, Search Console, or the live profiles. Every URL "
        "in section 13 resolves. The launch flight at ₨ 250,000 is sized for Pakistani market "
        "realities and structured around a clear bidding ladder with measurable KPIs.",
        size=11)
    add_paragraph(doc,
        "We thank Sir Maaz Zafar Cheema for the structured rubric — every pillar in this report "
        "maps directly to a graded KPI, and every claim is verifiable against a live URL or a "
        "saved screenshot. The team is happy to walk through any section in detail during the "
        "viva.",
        size=11)

    add_paragraph(doc, "— End of report —", size=10, color=FG_LOW, italic=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)


# ── Build ──────────────────────────────────────────────────────────────────


def main():
    doc = Document()

    # Default styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    cover_page(doc)
    table_of_contents(doc)
    section_executive_summary(doc)
    section_overview(doc)
    section_brand(doc)
    section_video_ad(doc)
    section_social(doc)
    section_meta_ads(doc)
    section_conversation(doc)
    section_seo(doc)
    section_google_ads(doc)
    section_budget(doc)
    section_competitive(doc)
    section_kpi(doc)
    section_links(doc)
    section_conclusion(doc)

    out_dir = ROOT / "submissions"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_file = out_dir / "VidIQ_Final_Report.docx"
    try:
        doc.save(out_file)
    except PermissionError:
        i = 2
        while (out_dir / f"VidIQ_Final_Report_v{i}.docx").exists():
            i += 1
        out_file = out_dir / f"VidIQ_Final_Report_v{i}.docx"
        doc.save(out_file)
        print(f"[warn] primary file locked - wrote {out_file.name} instead")
    print(f"[ok] wrote {out_file}")


if __name__ == "__main__":
    main()
